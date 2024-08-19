from flask import Flask, request, jsonify, render_template, session, redirect, url_for, make_response
from flask_pymongo import PyMongo, ObjectId
from werkzeug.security import generate_password_hash, check_password_hash
import re
from decorators.auth_required import auth_required
from bson import ObjectId
import requests
from forms.RegisterForm import RegisterForm
from forms.InvoiceForm import InvoiceForm
from docx import Document
import io


app = Flask(__name__)
app.config['MONGO_URI'] = 'mongodb://localhost:27017/invoices'
app.config['SECRET_KEY'] = 'abcd4321'

mongo = PyMongo(app)
db = mongo.db

@app.route('/register', methods = ['GET','POST'])

def register():
    form=RegisterForm()
    data = request.form
    email = data['email']
    password = data['password']
    firstname = data['firstname']

    if not re.fullmatch(r'[A-Za-z0-9@#$%^&+=]{8,}', password):
        return jsonify({
            'message':'Hasło musi składać się z minimum 8 znaków itp... bla bla'
        })


    existing_user = db.users.find_one({'email': email})

    

    if existing_user is None:

        hashed_password = generate_password_hash(password)
        new_data = {
            'email':email,
            'password':hashed_password,
            'firstname':firstname
        }
        user = db.users.insert_one(new_data)
        session['user_id'] = str(user.inserted_id)
        return redirect(url_for('dashboard'))
    
    else:
        error = "Uzupełnij wszystkie pola"
        return render_template('register.html', form=form)

@app.route("/login", methods = ['POST'])
def login():
    if request.method == "POST":
        data = request.form
        email = data['email']
        password = data['password']
        user = db['users'].find_one({'email':email})

        if user and check_password_hash(user['password'], password):
            session['user_id'] = str(user['_id'])
            return redirect(url_for('dashboard'))

        else:
            error = 'Podano błędne hasło lub email'
            return render_template('login.html', error=error)

@app.route('/logout')
def logout():
    session.pop('user_id')
    return redirect('/zaloguj-sie')




@app.route("/show-all")

@auth_required

def show_all():

    

        data = list(mongo.db['users'].find({},{'password':0}))
        all_users = mongo.db['users'].count_documents({})

        for item in data:
            item['_id'] = str(item['_id'])


        return jsonify({
        'data': data,
        'message': f'Pobrano listę wszystkich użytkowników {all_users}'
        })

@app.route('/user-details/<id>')
def user_details(id):
    user = db['users'].find_one({'_id': ObjectId(id)})
    user['_id'] = str(user['_id'])
    return jsonify({
          'data': user,
          'message':'Informacje o użytkowniku'
     })



@app.route("/whoami")
@auth_required
def whoami():
     user = db['users'].find_one({'_id': ObjectId(session['user_id'])})
     user['_id'] = str(user['_id'])
     return jsonify({
          'data': user,
          'message':'Informacje o użytkowniku'
     })
@app.route('/generate-docx/<id>')
def generate_docx(id):
     invoice = dict(db.invoices.find_one({"_id": ObjectId(id)}))
     docx = Document("doc/fv.docx")

     invoice['netto'] = round(invoice['netto'],2)
     invoice['brutto'] = round(invoice['brutto'],2)


    
     for p in docx.paragraphs:
        for key, value in invoice.items():
            if f"{{{{{key}}}}}" in p.text:
                p.text = p.text.replace(f"{{{{{key}}}}}", str(value))

     for table in docx.tables:
          for row in table.rows:
               for cell in row.cells:
                   for key, value in invoice.items():
                         if f"{{{{{key}}}}}" in cell.text:
                              cell.text = cell.text.replace(f"{{{{{key}}}}}", str(value))

     

     buffer = io.BytesIO()
     docx.save(buffer)
     # docx.save("doc/plik1.docx")

     response = make_response(buffer.getvalue())
     response.headers['Content-Disposition'] = f"attachment; filename=faktura-{invoice['invoice_number'].lower()}.docx"
     response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
     return response

@app.route('/delete-invoice/<id>')
def delete_invoice(id):
    try:
         deleted = db['invoices'].delete_one({'_id':ObjectId(id)})
         print(deleted)

    except:
        print("błąd")

    return redirect('/moje-faktury')

#Views


@app.route("/")
@auth_required
def dashboard():
     id = session["user_id"]
     user = requests.get(f"http://127.0.0.1:5000/user-details/{id}").json()
     print(user)
     return render_template("index.html", user=user)

@app.route('/zaloguj-sie', methods= ['GET','POST']) 
@auth_required
def login_page():
    return render_template('login.html')

@app.route('/zarejestruj-sie', methods= ['GET','POST'])
@auth_required 
def register_page():
    form = RegisterForm()

    if form.validate_on_submit():
        email = form.email.data
        firstname = form.firstname.data
        password = form.password.data

        existing_user = db.users.find_one({'email': email})

    

        if existing_user is None:

             hashed_password = generate_password_hash(password)
             new_data = {
                    'email':email,
                    'password':hashed_password,
                    'firstname':firstname
        }
             user = db.users.insert_one(new_data)
             session['user_id'] = str(user.inserted_id)
             return redirect(url_for('dashboard'))
    
        else:
             form.email.errors.append('Użytkownik z tym adresem email już istnieje')
    else:
        print(form.errors)
        
    return render_template('register.html', form=form)

@app.route('/moje-faktury', methods=['POST','GET'])
@auth_required
def invoices_page():

    invoices = list(db.invoices.find({'user_id': session['user_id']}))
   

    form = InvoiceForm()
    is_error = False
    is_success = False

    if request.method == 'POST':
        if form.validate_on_submit():
            form_data = form.data
            form_data['user_id'] = session['user_id']
            form_data['buyer_nip']= str(form_data['buyer_nip'])
            form_data['netto'] = float(form_data['item_price']) * float(form_data['item_quantity'])
            form_data['brutto'] = float(form_data['item_price']) * float(form_data['item_quantity']) * float('1.'+form_data['tax_value'])
            db.invoices.insert_one(form_data)
            is_success = True
            return redirect(url_for('invoices_page'))

           
        elif form.errors:
            is_error = True
        else:
            print(form.errors)

    return render_template('invoices.html',
                            form=form,
                            is_error=is_error,
                            is_success=is_success,
                            invoices=invoices)

@app.route('/moje-faktury/<id>')
@auth_required
def invoice_details_page(id):
    invoice = dict(db.invoices.find_one({"_id":ObjectId(id)}))

    return render_template("invoice_details.html", invoice=invoice)


@app.route('/edytuj-fakture/<id>', methods=['POST','GET'])
@auth_required
def edit_invoice(id):
    form = InvoiceForm()
    invoice = db['invoices'].find_one({'_id':ObjectId(id)})

    is_error = False
    is_success = False

    if request.method == 'POST':
        if form.validate_on_submit():
            form_data = form.data
            form_data['user_id'] = session['user_id']
            form_data['buyer_nip']= str(form_data['buyer_nip'])
            form_data['netto'] = float(form_data['item_price']) * float(form_data['item_quantity'])
            form_data['brutto'] = float(form_data['item_price']) * float(form_data['item_quantity']) * float('1.'+form_data['tax_value'])
            db['invoices'].update_one({'_id':ObjectId(id)}, {'$set':form_data}, upsert=False)
            print(id, 'test')
            
            is_success = True
            return redirect(url_for('invoices_page'))

           
    elif form.errors:
            is_error = True
    else:
            print(form.errors)


    return render_template('edit_invoice.html', form = form, invoice = invoice, is_error=is_error, is_success=is_success)